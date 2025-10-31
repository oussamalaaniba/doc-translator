# app.py — Doc Translator + Background image (local & cloud-safe)
import os, io, re, json, time, tempfile, subprocess, shutil, base64, zipfile
from io import BytesIO
from typing import List, Dict, Tuple, Any, Optional
from pathlib import Path

import streamlit as st
from dotenv import load_dotenv

# ===== Optional deps =====
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree as ET
except Exception:
    Document = None
    def qn(x): return x  # type: ignore
    ET = None

# PPTX (optional)
TRANSLATE_PPTX_AVAILABLE = True
try:
    from pptx_utils import translate_pptx_preserve_styles  # noqa
except Exception:
    TRANSLATE_PPTX_AVAILABLE = False
    translate_pptx_preserve_styles = None  # type: ignore

# =================== Config ===================
st.set_page_config(page_title="Doc Translator", page_icon="🌐", layout="centered")
load_dotenv()
os.makedirs("outputs", exist_ok=True)

for k, v in {"translated_bytes": None, "translated_name": None, "translated_mime": None, "last_filename": None}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# =================== Background (local or cloud) ===================
def set_background_auto(default_filename: str = "bg.png", darken: float = 0.35):
    """
    Utilise bg.png à côté de app.py si présent. Sinon:
    - st.secrets['BACKGROUND_URL']  : URL https
    - st.secrets['BACKGROUND_BASE64']: base64 d'une image (sans 'data:')

    'darken' applique un voile sombre pour la lisibilité (0..0.6).
    """
    # 1) image locale (même dossier que app.py)
    here = Path(__file__).parent
    local = here / default_filename
    img_b64 = None
    url = None

    if local.exists():
        try:
            img_b64 = base64.b64encode(local.read_bytes()).decode()
        except Exception:
            img_b64 = None

    # 2) secrets (pour le cloud)
    if img_b64 is None:
        url = (st.secrets.get("BACKGROUND_URL", None) if hasattr(st, "secrets") else None)
        if not url:
            b64 = (st.secrets.get("BACKGROUND_BASE64", None) if hasattr(st, "secrets") else None)
            if b64:
                img_b64 = b64.strip()

    if not img_b64 and not url:
        return  # pas de fond

    css_bg = (
        f'linear-gradient(rgba(0,0,0,{darken}), rgba(0,0,0,{darken})), '
        + (f'url("data:image/png;base64,{img_b64}")' if img_b64 else f'url("{url}")')
        + ' center center / cover no-repeat fixed'
    )

    st.markdown(
        f"""
        <style>
        [data-testid="stAppViewContainer"] {{
            background: {css_bg};
        }}
        [data-testid="stHeader"] {{
            background: rgba(0,0,0,0);
        }}
        .main .block-container {{
            background: rgba(255,255,255,0.08);
            border-radius: 12px;
            padding: 1rem 1.25rem;
        }}
        [data-testid="stSidebar"] > div:first-child {{
            background: rgba(255,255,255,0.75);
            backdrop-filter: blur(4px);
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )

# Active l'arrière-plan:
set_background_auto("bg.png", darken=0.35)

# =================== Secrets & env ===================
def _get_secret(name: str, default=None):
    try:
        return st.secrets[name]
    except Exception:
        return os.getenv(name, default)

def get_openai_key() -> Optional[str]:
    return _get_secret("OPENAI_API_KEY") or _get_secret("openai_api_key") or _get_secret("OPENAI_KEY")

def has_ocr_binary() -> bool:
    return shutil.which("ocrmypdf") is not None

def is_cloud_environment() -> bool:
    return (str(_get_secret("DISABLE_OCR", "0")) == "1") or (not has_ocr_binary())

OCR_AVAILABLE_LOCALLY = has_ocr_binary()
RUNNING_IN_CLOUD = is_cloud_environment()
SHOW_OCR_BUTTON = OCR_AVAILABLE_LOCALLY and not RUNNING_IN_CLOUD

# =================== IO util ===================
def save_output_file(file_bytes: bytes, file_name: str) -> str:
    path = os.path.join("outputs", file_name)
    with open(path, "wb") as f:
        f.write(file_bytes)
    return path

# =================== OpenAI translate (robust batch) ===================
def _clean_json_content(content: str) -> str:
    s = (content or "").strip()
    if s.startswith("```"):
        s = s.strip("`")
        s = re.sub(r"^json\n", "", s, flags=re.IGNORECASE)
    return s

def translate_batch(texts: List[str], src: str = "fr", tgt: str = "en",
                    *, timeout: int = 60, max_retries: int = 3) -> List[str]:
    """
    Batch robuste: indexe chaque item, reconstruit par index, complète manquants one-by-one.
    """
    if not texts:
        return []
    api_key = get_openai_key()
    if not api_key:
        return texts

    inputs = [("" if t is None else str(t)) for t in texts]
    n = len(inputs)

    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)

        payload_items = [{"i": i, "t": inputs[i]} for i in range(n)]
        system = (
            "You are a professional translator. Translate EACH item from 'src' to 'tgt'. "
            "Preserve meaning, register, numbers, punctuation. "
            "Do NOT translate tokens like [[TOK#]], [[GLOS#]], [[KEEP#]] nor URLs/emails. "
            "Return STRICT JSON with EXACTLY N items and SAME indices.\n"
            'Schema: {\"items\":[{\"i\":0,\"t\":\"...\"}, ...]}\n'
            "If an item is empty, return empty string. No commentary."
        )
        user_payload = {"src": src, "tgt": tgt, "N": n, "items": payload_items}
        user = json.dumps(user_payload, ensure_ascii=False)

        last_err = None
        for attempt in range(max_retries):
            try:
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
                    temperature=0,
                    response_format={"type": "json_object"},
                )
                content = _clean_json_content(resp.choices[0].message.content or "")
                obj = json.loads(content)

                out = [""] * n
                got = set()
                items = obj.get("items", []) if isinstance(obj, dict) else obj
                if isinstance(items, list):
                    for it in items:
                        try:
                            i = int(it.get("i", -1)); t = it.get("t", "")
                        except Exception:
                            continue
                        if 0 <= i < n:
                            out[i] = ("" if t is None else str(t)).replace("\u00A0", " ")
                            got.add(i)

                if len(got) == n:
                    return out

                missing = [i for i in range(n) if i not in got]
                st.warning(f"Bulk translation returned {len(got)}/{n}. Retrying {len(missing)} missing item(s) one-by-one.")
                for i in missing:
                    one = translate_batch([inputs[i]], src, tgt, timeout=timeout, max_retries=1)
                    out[i] = (one[0] if one else inputs[i])
                    time.sleep(0.05)
                return out

            except Exception as e:
                last_err = e
                time.sleep(1.2 * (attempt + 1))

        st.warning(f"Bulk translation failed ({last_err}); falling back to one-by-one.")
        out: List[str] = []
        for t in inputs:
            out.extend(translate_batch([t], src, tgt, timeout=timeout, max_retries=1))
            time.sleep(0.05)
        return out

    except Exception as e:
        st.error(f"Erreur API de traduction : {e}")
        return texts

# =================== DOCX helpers & pipeline ===================
# WordprocessingML + DrawingML namespaces
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",      # SmartArt
    "dgm2010": "http://schemas.microsoft.com/office/drawing/2010/diagram",  # SmartArt 2010+
}

ACRONYM_REGEX = r"\b[A-Z]{3,}\b"

def _normalize_text_after(t: str) -> str:
    t = t.replace("\u00A0", " ")
    t = re.sub(r"[ \t]{2,}", " ", t)
    return t

def _make_token(prefix: str, idx: int) -> str:
    return f"[[{prefix}{idx}]]"

def _protect_patterns(text: str) -> Tuple[str, Dict[str, str]]:
    mapping: Dict[str, str] = {}
    idx = 0
    patterns = [r"https?://\S+", r"[\w\.-]+@[\w\.-]+\.\w+", r"\{[^{}]+\}", r"%[sdif]", ACRONYM_REGEX]
    def repl(m):
        nonlocal idx
        val = m.group(0); key = _make_token("TOK", idx); idx += 1
        mapping[key] = val; return key
    for pat in patterns:
        text = re.sub(pat, repl, text)
    return text, mapping

def _protect_terms_ci(text: str, terms: List[str], prefix: str) -> Tuple[str, Dict[str, str]]:
    mapping: Dict[str, str] = {}
    if not terms: return text, mapping
    terms_sorted = sorted(terms, key=lambda s: len(s), reverse=True)
    idx = 0
    for term in terms_sorted:
        pat = rf"(?i){re.escape(term)}"
        def repl(m):
            nonlocal idx
            key = _make_token(prefix, idx); idx += 1
            mapping[key] = m.group(0); return key
        text = re.sub(pat, repl, text)
    return text, mapping

def _protect_glossary_ci(text: str, glossary: Dict[str, str]) -> Tuple[str, Dict[str, str]]:
    if not glossary: return text, {}
    items = sorted(glossary.items(), key=lambda kv: len(kv[0]), reverse=True)
    mapping: Dict[str, str] = {}
    idx = 0
    for src, tgt in items:
        pat = rf"(?i){re.escape(src)}"
        def repl(m):
            nonlocal idx
            key = _make_token("GLOS", idx); idx += 1
            mapping[key] = tgt; return key
        text = re.sub(pat, repl, text)
    return text, mapping

def _parse_glossary_csv(csv_text: str) -> Dict[str, str]:
    d: Dict[str, str] = {}
    if not csv_text: return d
    for line in csv_text.splitlines():
        line = line.strip()
        if not line or line.startswith("#"): continue
        parts = [p.strip() for p in line.split(",")]
        if len(parts) >= 2 and parts[0] and parts[1]: d[parts[0]] = parts[1]
    return d

def _parse_dnt_terms(text: str) -> List[str]:
    if not text: return []
    items: List[str] = []
    for chunk in re.split(r"[\n,]", text):
        t = chunk.strip()
        if t: items.append(t)
    seen = set(); out: List[str] = []
    for t in items:
        k = t.lower()
        if k not in seen:
            seen.add(k); out.append(t)
    return out

def _preprocess_text_for_translation(t: str, glossary: Dict[str, str], dnt_terms: List[str]) -> Tuple[str, Dict[str, Dict[str, str]]]:
    t1, map_tok  = _protect_patterns(t)
    t2, map_keep = _protect_terms_ci(t1, dnt_terms, "KEEP")
    t3, map_glos = _protect_glossary_ci(t2, glossary)
    return t3, {"GLOS": map_glos, "KEEP": map_keep, "TOK": map_tok}

def _postprocess_translation(t: str, m: Dict[str, Dict[str, str]]) -> str:
    for prefix in ("GLOS", "KEEP", "TOK"):
        for token, val in m.get(prefix, {}).items():
            t = t.replace(token, val)
    return _normalize_text_after(t)

# --- paragraph writers ---
def _set_paragraph_text_preserve_para_style(p, new_text: str) -> None:
    """
    Écrit le texte sans hériter de la mise en forme du 1er run (évite 'tout en gras').
    """
    if getattr(p, "runs", None) and p.runs:
        for r in p.runs: r.text = ""
        p.add_run(new_text)  # run neutre
    else:
        p.add_run(new_text)

def _has_mixed_bold_runs(p) -> bool:
    try:
        runs = [r for r in getattr(p, "runs", []) if (r.text or "").strip()]
        if len(runs) < 2: return False
        vals = []
        for r in runs:
            b = r.bold
            if b is None: b = False
            vals.append(bool(b))
        return any(vals) and not all(vals)
    except Exception:
        return False

# --- XML/DML traversal ---
def _collect_xml_paragraphs(root):
    """<w:p> dans w:txbxContent, v:textbox et w:sdtContent."""
    if root is None or ET is None: return []
    return list(root.xpath(
        './/w:txbxContent//w:p | .//v:textbox//w:p | .//w:sdtContent//w:p',
        namespaces=NSMAP
    ))

def _collect_dml_paragraphs(root):
    """Paragraphes DrawingML (a:p) dans formes/WordArt: w:drawing//a:txBody//a:p"""
    if root is None or ET is None: return []
    return list(root.xpath('.//w:drawing//a:txBody//a:p', namespaces=NSMAP))

def _get_wt_texts(p_xml):
    if ET is None or p_xml is None: return []
    return p_xml.xpath('.//w:t', namespaces=NSMAP)

def _get_at_texts(a_p):
    if ET is None or a_p is None: return []
    return a_p.xpath('.//a:t', namespaces=NSMAP)

def _get_xml_para_text(p_xml) -> str:
    ts = _get_wt_texts(p_xml)
    return "".join((t.text or "") for t in ts) if ts else ""

def _set_xml_paragraph_text(p_xml, new_text: str) -> None:
    ts = _get_wt_texts(p_xml)
    if ts:
        ts[0].text = new_text
        for t in ts[1:]: t.text = ''
    else:
        r = ET.Element(qn('w:r')); t = ET.SubElement(r, qn('w:t')); t.text = new_text
        p_xml.append(r)

def _get_dml_para_text(a_p) -> str:
    ts = _get_at_texts(a_p)
    return "".join((t.text or "") for t in ts) if ts else ""

def _set_dml_para_text(a_p, new_text: str) -> None:
    ts = _get_at_texts(a_p)
    if ts:
        ts[0].text = new_text
        for t in ts[1:]: t.text = ''
    else:
        r = ET.Element(qn('a:r')); t = ET.SubElement(r, qn('a:t')); t.text = new_text
        a_p.append(r)

# --- Detect & ignore TOC ---
def _looks_like_toc_line(txt: str) -> bool:
    if not txt: return False
    s = txt.strip()
    if ("\t" in s or "|" in s) and re.search(r'\d+\s*$', s): return True
    if re.search(r'\.{3,}\s*\d+\s*$', s): return True  # ..... 12
    if len(s) <= 120 and re.search(r'\d+\s*$', s) and re.search(r'\b\d', s): return True
    return False

def _is_toc_para_docx(p) -> bool:
    try:
        el = p._element
        if el.xpath('.//w:instrText[contains(., "TOC")]', namespaces=NSMAP): return True
        for fld in el.xpath('.//w:fldSimple', namespaces=NSMAP):
            instr = fld.get(qn('w:instr'))
            if instr and 'TOC' in instr: return True
        sname = (getattr(p, "style", None) and p.style.name or "").lower()
        if sname.startswith("toc"): return True
        if _looks_like_toc_line(getattr(p, "text", "") or ""): return True
    except Exception:
        pass
    return False

def _is_toc_para_xml(p_xml) -> bool:
    try:
        if p_xml.xpath('.//w:instrText[contains(., "TOC")]', namespaces=NSMAP): return True
        for fld in p_xml.xpath('.//w:fldSimple', namespaces=NSMAP):
            instr = fld.get(qn('w:instr'))
            if instr and 'TOC' in instr: return True
        txt_nodes = p_xml.xpath('.//w:t', namespaces=NSMAP)
        txt = ''.join((n.text or '') for n in txt_nodes).strip() if txt_nodes else ''
        if _looks_like_toc_line(txt): return True
    except Exception:
        pass
    return False

def _is_toc_para_dml(a_p) -> bool:
    try:
        ts = _get_at_texts(a_p)
        txt = ''.join((t.text or '') for t in ts) if ts else ''
        if _looks_like_toc_line(txt): return True
        if re.match(r'(?i)^\s*(sommaire|summary|table\s+of\s+contents)\s*$', txt): return True
    except Exception:
        pass
    return False

# --- Collect all paragraphs ---
def _collect_paragraph_objects_from_doc(doc) -> List[Tuple[str, Any]]:
    """
    Returns [(kind, ref)] where:
      kind="docx" → python-docx Paragraph
      kind="xml"  → lxml <w:p> (text boxes / content controls)
      kind="dml"  → lxml <a:p> (DrawingML in shapes/WordArt)
    """
    out: List[Tuple[str, Any]] = []
    seen = set()

    # Body
    for p in doc.paragraphs:
        if id(p) not in seen:
            seen.add(id(p)); out.append(("docx", p))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if id(p) not in seen:
                        seen.add(id(p)); out.append(("docx", p))

    try:
        body_xml = doc.element.body
        for p_xml in _collect_xml_paragraphs(body_xml):
            out.append(("xml", p_xml))
        for a_p in _collect_dml_paragraphs(body_xml):
            out.append(("dml", a_p))
    except Exception:
        pass

    # Headers/Footers (incluant variantes)
    for section in doc.sections:
        for attr in ("header", "first_page_header", "even_page_header"):
            part = getattr(section, attr, None)
            if not part: continue
            for p in part.paragraphs:
                if id(p) not in seen:
                    seen.add(id(p)); out.append(("docx", p))
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if id(p) not in seen:
                                seen.add(id(p)); out.append(("docx", p))
            try:
                root = part._element
                for p_xml in _collect_xml_paragraphs(root):
                    out.append(("xml", p_xml))
                for a_p in _collect_dml_paragraphs(root):
                    out.append(("dml", a_p))
            except Exception:
                pass

        for attr in ("footer", "first_page_footer", "even_page_footer"):
            part = getattr(section, attr, None)
            if not part: continue
            for p in part.paragraphs:
                if id(p) not in seen:
                    seen.add(id(p)); out.append(("docx", p))
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if id(p) not in seen:
                                seen.add(id(p)); out.append(("docx", p))
            try:
                root = part._element
                for p_xml in _collect_xml_paragraphs(root):
                    out.append(("xml", p_xml))
                for a_p in _collect_dml_paragraphs(root):
                    out.append(("dml", a_p))
            except Exception:
                pass

    return out

# --- SmartArt translator (DOCX parts in word/diagrams) ---
def _translate_docx_smartart_parts(docx_bytes: bytes, src: str, tgt: str,
                                   glossary: Dict[str, str], dnt_terms: List[str]) -> bytes:
    """
    Ouvre le DOCX comme un ZIP, modifie word/diagrams/data*.xml :
    - extrait tous les <a:t> (et <dgm:t> fallback) présents dans le modèle SmartArt
    - applique le même pré/post-traitement (glossaire, DNT, tokens)
    - remplace le texte par la traduction
    """
    if ET is None:
        return docx_bytes

    bio_in = BytesIO(docx_bytes)
    with zipfile.ZipFile(bio_in, "r") as zin:
        bio_out = BytesIO()
        with zipfile.ZipFile(bio_out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            modified = False

            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename.startswith("word/diagrams/") and item.filename.endswith(".xml"):
                    try:
                        root = ET.fromstring(data)
                    except Exception:
                        # XML non lisible → on copie tel quel
                        zout.writestr(item, data)
                        continue

                    # Tous les nœuds de texte dans le modèle SmartArt
                    nodes = root.xpath(".//a:t | .//dgm:t | .//dgm2010:t", namespaces=NSMAP)
                    if not nodes:
                        zout.writestr(item, data)
                        continue

                    # Pré-traitement + traduction batch
                    pre_list, maps_list = [], []
                    for n in nodes:
                        txt = n.text or ""
                        pre, maps = _preprocess_text_for_translation(txt, glossary, dnt_terms)
                        pre_list.append(pre)
                        maps_list.append(maps)

                    tr_list = translate_batch(pre_list, src=src, tgt=tgt)

                    # Post-traitement + réécriture
                    for n, tr, maps in zip(nodes, tr_list, maps_list):
                        n.text = _postprocess_translation(tr, maps)

                    new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                    zout.writestr(item, new_xml)
                    modified = True
                else:
                    # Parts non concernées → copie inchangée
                    zout.writestr(item, data)

        return bio_out.getvalue() if modified else docx_bytes

# --- DOCX pipeline ---
def translate_docx_preserve_styles(src_bytes: bytes, src: str = "fr", tgt: str = "en") -> bytes:
    if Document is None:
        raise RuntimeError("python-docx manquant : impossible de traiter un DOCX.")

    GLOSSARY = _parse_glossary_csv(st.session_state.get("glossary_csv", ""))
    DNT      = _parse_dnt_terms(st.session_state.get("dnt_terms", ""))

    doc = Document(BytesIO(src_bytes))
    collected = _collect_paragraph_objects_from_doc(doc)

    refs: List[Tuple[str, Any]] = []
    to_translate: List[str] = []
    maps_list: List[Dict[str, Dict[str, str]]] = []

    inside_toc = False
    for kind, ref in collected:
        # Détection bloc TOC (titre Sommaire/Summary)
        para_text = ""
        if kind == "docx":
            para_text = getattr(ref, "text", "") or ""
            if re.match(r'(?i)^\s*(sommaire|summary|table\s+of\s+contents)\s*$', para_text.strip()):
                inside_toc = True; continue
            try: sname = (ref.style.name or "").lower()
            except Exception: sname = ""
            if inside_toc and (("heading" in sname) or ("titre" in sname) or re.match(r'(?i)^\s*(introduction|1\s)', para_text.strip())):
                inside_toc = False
        elif kind == "dml":
            try:
                ts = _get_at_texts(ref)
                para_text = ''.join((t.text or '') for t in ts) if ts else ''
                if re.match(r'(?i)^\s*(sommaire|summary|table\s+of\s+contents)\s*$', para_text.strip()):
                    inside_toc = True; continue
                if inside_toc and re.match(r'(?i)^\s*(introduction|1\s)', para_text.strip()):
                    inside_toc = False
            except Exception:
                pass

        if inside_toc:
            continue
        if (kind == "docx" and _is_toc_para_docx(ref)) or (kind == "xml" and _is_toc_para_xml(ref)) or (kind == "dml" and _is_toc_para_dml(ref)):
            continue

        # Préservation gras: run-par-run si mixé
        if kind == "docx" and _has_mixed_bold_runs(ref):
            for idx, run in enumerate(ref.runs):
                txt = (run.text or '')
                if not txt.strip(): continue
                pre, maps = _preprocess_text_for_translation(txt, GLOSSARY, DNT)
                refs.append(("docx_run", (ref, idx))); to_translate.append(pre); maps_list.append(maps)
            continue

        if kind == "docx":
            original = getattr(ref, "text", "") or ""
        elif kind == "xml":
            original = _get_xml_para_text(ref)
        else:
            original = _get_dml_para_text(ref)

        if original.strip():
            pre, maps = _preprocess_text_for_translation(original, GLOSSARY, DNT)
            refs.append((kind, ref)); to_translate.append(pre); maps_list.append(maps)

    if not to_translate:
        # Aucun paragraphe standard, mais on pourra tout de même traiter SmartArt
        out_bytes = src_bytes
        out_bytes = _translate_docx_smartart_parts(out_bytes, src=src, tgt=tgt, glossary=GLOSSARY, dnt_terms=DNT)
        return out_bytes

    translated_all: List[str] = []
    CHUNK = 16
    total = len(to_translate)
    prog = st.progress(0.0)
    for i in range(0, total, CHUNK):
        chunk = to_translate[i:i+CHUNK]
        out = translate_batch(chunk, src, tgt, timeout=60)
        if len(out) < len(chunk):
            out = out + chunk[len(out):]
        translated_all.extend(out)
        prog.progress(min(1.0, len(translated_all)/total))
    prog.empty()

    for (kind, ref), tr, maps in zip(refs, translated_all, maps_list):
        final_text = _postprocess_translation(tr, maps)
        if kind == "docx":
            _set_paragraph_text_preserve_para_style(ref, final_text)
        elif kind == "docx_run":
            para, idx = ref
            try:
                para.runs[idx].text = final_text
            except Exception:
                para.add_run(final_text)
        elif kind == "xml":
            _set_xml_paragraph_text(ref, final_text)
        else:
            _set_dml_para_text(ref, final_text)

    bio = BytesIO(); doc.save(bio); bio.seek(0)
    out_bytes = bio.read()

    # ⬇️ Ajout : passe sur les parts SmartArt (word/diagrams/data*.xml)
    out_bytes = _translate_docx_smartart_parts(
        out_bytes, src=src, tgt=tgt, glossary=GLOSSARY, dnt_terms=DNT
    )

    return out_bytes

# =================== PDF ===================
def pdf_has_text(src_bytes: bytes, min_chars: int = 20) -> bool:
    if fitz is None: return False
    try:
        doc = fitz.open(stream=src_bytes, filetype="pdf")
        has = False
        for page in doc:
            blocks = page.get_text("blocks")
            for b in blocks:
                if len(b) >= 5 and isinstance(b[4], str) and len(b[4].strip()) >= min_chars:
                    has = True; break
            if has: break
        doc.close(); return has
    except Exception:
        return False

def ocr_pdf_with_ocrmypdf(src_bytes: bytes, lang: str = "fra") -> bytes:
    if RUNNING_IN_CLOUD: return src_bytes
    try:
        if pdf_has_text(src_bytes): return src_bytes
        with tempfile.TemporaryDirectory() as td:
            inp = os.path.join(td, "in.pdf"); outp = os.path.join(td, "out.pdf")
            with open(inp, "wb") as f: f.write(src_bytes)
            cmd = ["ocrmypdf", "--skip-text", f"--language={lang}", "--output-type", "pdf", "--optimize", "0", "--fast-web-view", "0", inp, outp]
            subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=120)
            with open(outp, "rb") as f: return f.read()
    except subprocess.TimeoutExpired:
        st.warning("⏱️ OCR trop long → on continue sans OCR.")
    except Exception as e:
        st.warning(f"OCR ignoré ({e})")
    return src_bytes

def translate_pdf_overlay(src_bytes: bytes, src: str = "fr", tgt: str = "en") -> bytes:
    if fitz is None:
        raise RuntimeError("PyMuPDF manquant : traduction PDF indisponible.")
    doc = fitz.open(stream=src_bytes, filetype="pdf")
    for page in doc:
        blocks = page.get_text("blocks")
        texts = [b[4] for b in blocks if len(b) >= 5 and isinstance(b[4], str) and b[4].strip()]
        if not texts: continue
        translated = translate_batch(texts, src, tgt)
        for (x0, y0, x1, y1, _txt, *_) in blocks:
            rect = fitz.Rect(x0, y0, x1, y1)
            page.add_redact_annot(rect, fill=(1, 1, 1))
        page.apply_redactions()
        def insert_text_fit(page, rect, text, fontname="Helvetica", max_size=11, min_size=6, step=0.5, align=0):
            size = max_size
            while size >= min_size:
                used = page.insert_textbox(rect, text, fontname=fontname, fontsize=size, align=align)
                if used >= 0: return True
                page.add_redact_annot(rect, fill=(1, 1, 1)); page.apply_redactions()
                size -= step
            page.insert_textbox(rect, text, fontname=fontname, fontsize=min_size, align=align); return False
        for (x0, y0, x1, y1, _txt, *_), new_text in zip(blocks, translated):
            rect = fitz.Rect(x0, y0, x1, y1)
            insert_text_fit(page, rect, new_text, max_size=11, min_size=6, step=0.5, align=0)
    out = io.BytesIO(); doc.save(out); doc.close(); out.seek(0)
    return out.read()

# =================== UI ===================
st.title("🌐 Document Translator ")

src_lang = st.selectbox("Langue source", ["fr", "en", "es", "de"], index=0)
tgt_lang = st.selectbox("Langue cible", ["en", "fr", "es", "de"], index=1)

with st.expander("⚙️ Options traduction (DOCX)"):
    st.session_state["glossary_csv"] = st.text_area(
        "Glossaire source,target (CSV, une paire par ligne)",
        value=st.session_state.get("glossary_csv", ""),
        placeholder="serveur,server\nclient,customer",
    )
    st.session_state["dnt_terms"] = st.text_area(
        "Termes à NE PAS traduire (un par ligne ou séparés par des virgules)",
        value=st.session_state.get("dnt_terms", ""),
        placeholder="OpenAI\nGPU\nGPT-4o",
    )
    st.caption("Astuce : le glossaire force une traduction précise. Les DNT seront laissés tels quels.")

uploaded = st.file_uploader("Dépose ton fichier .docx, .pptx ou .pdf", type=["docx", "pptx", "pdf"])

if uploaded:
    data = uploaded.getvalue()
    st.info(f"Fichier reçu : {uploaded.name} ({len(data)} octets)")

    if st.session_state.get("last_filename") != uploaded.name:
        st.session_state.translated_bytes = None
        st.session_state.translated_name = None
        st.session_state.translated_mime = None
        st.session_state.last_filename = uploaded.name

    name_lower = uploaded.name.lower()

    # ======== PDF ========
    if name_lower.endswith(".pdf"):
        if SHOW_OCR_BUTTON:
            if st.button("1) OCR (si scanné) → 2) Traduire PDF", key="btn_translate_pdf_ocr"):
                with st.spinner("Traitement PDF (OCR si besoin + traduction)…"):
                    try:
                        lang_ocr = "fra" if src_lang == "fr" else src_lang
                        ocred = ocr_pdf_with_ocrmypdf(data, lang=lang_ocr)
                        translated = translate_pdf_overlay(ocred, src=src_lang, tgt=tgt_lang)
                        output_name = uploaded.name.replace(".pdf", f"_{tgt_lang}.pdf")
                        st.session_state.translated_bytes = translated
                        st.session_state.translated_name = output_name
                        st.session_state.translated_mime = "application/pdf"
                        save_path = save_output_file(translated, output_name)
                        st.success("✅ PDF traduit. Le bouton de téléchargement est prêt ci-dessous 👇")
                        st.info(f"💾 Fichier enregistré : {save_path}")
                    except Exception as e:
                        st.error(f"Erreur PDF/OCR: {e}")

            if st.button("Traduire PDF (sans OCR)", key="btn_translate_pdf_plain"):
                with st.spinner("Traduction PDF (sans OCR)…"):
                    try:
                        translated = translate_pdf_overlay(data, src=src_lang, tgt=tgt_lang)
                        output_name = uploaded.name.replace(".pdf", f"_{tgt_lang}.pdf")
                        st.session_state.translated_bytes = translated
                        st.session_state.translated_name = output_name
                        st.session_state.translated_mime = "application/pdf"
                        save_path = save_output_file(translated, output_name)
                        st.success("✅ PDF traduit (sans OCR).")
                        st.info(f"💾 Fichier enregistré : {save_path}")
                    except Exception as e:
                        st.error(f"Erreur PDF: {e}")
        else:
            st.warning("☁️ OCR désactivé en cloud. Traduction possible uniquement si le PDF est textuel.")
            if st.button("Traduire PDF (sans OCR)", key="btn_translate_pdf_cloud"):
                with st.spinner("Traduction PDF (sans OCR)…"):
                    try:
                        translated = translate_pdf_overlay(data, src=src_lang, tgt=tgt_lang)
                        output_name = uploaded.name.replace(".pdf", f"_{tgt_lang}.pdf")
                        st.session_state.translated_bytes = translated
                        st.session_state.translated_name = output_name
                        st.session_state.translated_mime = "application/pdf"
                        save_path = save_output_file(translated, output_name)
                        st.success("✅ PDF traduit (si textuel).")
                        st.info(f"💾 Fichier enregistré : {save_path}")
                    except Exception as e:
                        st.error(f"Erreur PDF: {e}")

    # ======== DOCX ========
    elif name_lower.endswith(".docx"):
        if st.button("Traduire DOCX", key="btn_translate_docx"):
            with st.spinner("Traduction du DOCX en cours…"):
                try:
                    translated = translate_docx_preserve_styles(data, src=src_lang, tgt=tgt_lang)
                    output_name = uploaded.name.replace(".docx", f"_{tgt_lang}.docx")
                    st.session_state.translated_bytes = translated
                    st.session_state.translated_name = output_name
                    st.session_state.translated_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    save_path = save_output_file(translated, output_name)
                    st.success("✅ DOCX traduit. Le bouton de téléchargement est prêt ci-dessous 👇")
                    st.info(f"💾 Fichier enregistré : {save_path}")
                except Exception as e:
                    st.error(f"Erreur DOCX: {e}")

    # ======== PPTX ========
    elif name_lower.endswith(".pptx"):
        if not TRANSLATE_PPTX_AVAILABLE:
            st.error("Le module PPTX n'est pas disponible (pptx_utils manquant).")
        else:
            if st.button("Traduire PPTX", key="btn_translate_pptx"):
                with st.spinner("Traduction du PPTX en cours…"):
                    try:
                        translated = translate_pptx_preserve_styles(
                            data, src=src_lang, tgt=tgt_lang, translate_callable=translate_batch
                        )
                        output_name = uploaded.name.replace(".pptx", f"_{tgt_lang}.pptx")
                        st.session_state.translated_bytes = translated
                        st.session_state.translated_name = output_name
                        st.session_state.translated_mime = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
                        save_path = save_output_file(translated, output_name)
                        st.success("✅ PPTX traduit. Le bouton de téléchargement est prêt ci-dessous 👇")
                        st.info(f"💾 Fichier enregistré : {save_path}")
                    except Exception as e:
                        st.error(f"Erreur PPTX: {e}")

# Download button (common)
if st.session_state.translated_bytes:
    st.download_button(
        "⬇️ Télécharger le fichier traduit",
        data=st.session_state.translated_bytes,
        file_name=st.session_state.translated_name or "translated_file",
        mime=st.session_state.translated_mime or "application/octet-stream",
        key="download_translated_v1",
    )

st.divider()
st.write("⚙️ Conseils :")
st.write("- PDF : en Cloud, OCR désactivé. Les PDF scannés doivent être traités en local.")
